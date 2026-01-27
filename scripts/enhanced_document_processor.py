"""
Enhanced Document Processing with Advanced Structure Handling
Implements sophisticated document processing with structure preservation and rich metadata extraction.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import json

try:
    from langchain_community.document_loaders import (
        PyPDFLoader,
        Docx2txtLoader,
        UnstructuredPDFLoader,
        UnstructuredWordDocumentLoader,
        UnstructuredMarkdownLoader,
        UnstructuredHTMLLoader
    )
    from langchain.text_splitter import (
        RecursiveCharacterTextSplitter,
        MarkdownHeaderTextSplitter,
        HTMLHeaderTextSplitter
    )
    from langchain_core.documents import Document
    import pandas as pd
    from bs4 import BeautifulSoup
    import pdfplumber
    import docx
except ImportError:
    print("⚠️  Required packages not installed. Run: pip install langchain langchain-community pandas beautifulsoup4 pdfplumber python-docx")

@dataclass
class DocumentSection:
    """Represents a section in a document with hierarchical structure"""
    content: str
    heading: Optional[str] = None
    level: int = 0
    parent: Optional['DocumentSection'] = None
    children: List['DocumentSection'] = None
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.children is None:
            self.children = []
        if self.metadata is None:
            self.metadata = {}

class EnhancedDocumentProcessor:
    """Advanced document processor with structure preservation and rich metadata extraction"""
    
    def __init__(self, 
                 chunk_size: int = 1000, 
                 chunk_overlap: int = 200,
                 preserve_structure: bool = True):
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.preserve_structure = preserve_structure
        
        # Initialize text splitters
        self._init_text_splitters()
        
        print(f"✅ Enhanced Document Processor initialized with structure preservation")
    
    def _init_text_splitters(self):
        """Initialize various text splitters for different document types"""
        # General text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            keep_separator=True,
        )
        
        # Markdown-aware splitter
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "header_1"),
                ("##", "header_2"),
                ("###", "header_3"),
                ("####", "header_4"),
            ]
        )
        
        # HTML-aware splitter
        self.html_splitter = RecursiveCharacterTextSplitter.from_language(
            language="html",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

    def process_file(self, file_path: str) -> List[Document]:
        """Process a file with structure preservation"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._process_pdf_enhanced(file_path)
        elif ext in ['.docx', '.doc']:
            return self._process_docx_enhanced(file_path)
        elif ext == '.md':
            return self._process_markdown(file_path)
        elif ext == '.html':
            return self._process_html(file_path)
        elif ext == '.txt':
            return self._process_text_enhanced(file_path)
        else:
            print(f"⚠️  Unsupported file type: {ext}")
            return []

    def _process_pdf_enhanced(self, file_path: str) -> List[Document]:
        """Process PDF with enhanced structure extraction"""
        try:
            documents = []
            
            # Extract text with layout preservation
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    
                    # Extract tables
                    tables = page.extract_tables()
                    tables_data = []
                    for table in tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables_data.append({
                            'table_content': df.to_dict(orient='records'),
                            'location': {'page': page_num}
                        })
                    
                    # Create document with rich metadata
                    doc = Document(
                        page_content=text,
                        metadata={
                            'source': Path(file_path).name,
                            'file_path': file_path,
                            'type': 'pdf',
                            'page': page_num,
                            'tables': tables_data,
                            'layout': {
                                'width': page.width,
                                'height': page.height,
                                'has_tables': bool(tables),
                                'has_images': bool(page.images)
                            }
                        }
                    )
                    documents.append(doc)
            
            # Split while preserving structure
            if self.preserve_structure:
                split_docs = self._split_with_structure(documents)
            else:
                split_docs = self.text_splitter.split_documents(documents)
            
            print(f"✓ Processed PDF: {Path(file_path).name} ({len(split_docs)} chunks)")
            return split_docs
            
        except Exception as e:
            print(f"✗ Error processing PDF {file_path}: {str(e)}")
            return []

    def _process_docx_enhanced(self, file_path: str) -> List[Document]:
        """Process DOCX with enhanced structure extraction"""
        try:
            doc = docx.Document(file_path)
            documents = []
            current_section = None
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = docx.Document(file_path)._body.add_paragraph()
                    text = paragraph.text
                    
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name[-1])
                        current_section = DocumentSection(
                            content=text,
                            heading=text,
                            level=level,
                            metadata={'type': 'heading', 'level': level}
                        )
                    else:
                        if current_section:
                            current_section.content += f"\n{text}"
                        else:
                            current_section = DocumentSection(
                                content=text,
                                metadata={'type': 'paragraph'}
                            )
                
                elif element.tag.endswith('tbl'):  # Table
                    table = docx.Document(file_path)._body.add_table(rows=1, cols=1)
                    table_data = []
                    
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    if current_section:
                        current_section.metadata['tables'] = current_section.metadata.get('tables', [])
                        current_section.metadata['tables'].append(table_data)
            
            # Convert sections to documents
            for section in self._flatten_sections([current_section]):
                doc = Document(
                    page_content=section.content,
                    metadata={
                        'source': Path(file_path).name,
                        'file_path': file_path,
                        'type': 'docx',
                        'structure': {
                            'heading': section.heading,
                            'level': section.level,
                        },
                        **section.metadata
                    }
                )
                documents.append(doc)
            
            # Split while preserving structure
            if self.preserve_structure:
                split_docs = self._split_with_structure(documents)
            else:
                split_docs = self.text_splitter.split_documents(documents)
            
            print(f"✓ Processed DOCX: {Path(file_path).name} ({len(split_docs)} chunks)")
            return split_docs
            
        except Exception as e:
            print(f"✗ Error processing DOCX {file_path}: {str(e)}")
            return []

    def _split_with_structure(self, documents: List[Document]) -> List[Document]:
        """Split documents while preserving their structure"""
        split_docs = []
        
        for doc in documents:
            # Extract any existing structure information
            structure = doc.metadata.get('structure', {})
            
            # Split the content
            splits = self.text_splitter.split_text(doc.page_content)
            
            # Create new documents with preserved structure
            for i, split in enumerate(splits):
                new_doc = Document(
                    page_content=split,
                    metadata={
                        **doc.metadata,
                        'chunk_info': {
                            'chunk_index': i,
                            'total_chunks': len(splits)
                        }
                    }
                )
                split_docs.append(new_doc)
        
        return split_docs

    def _flatten_sections(self, sections: List[DocumentSection]) -> List[DocumentSection]:
        """Flatten nested sections into a list"""
        flattened = []
        for section in sections:
            flattened.append(section)
            if section.children:
                flattened.extend(self._flatten_sections(section.children))
        return flattened

    def extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract rich metadata from a document"""
        metadata = {
            'content_stats': {
                'length': len(doc.page_content),
                'num_paragraphs': doc.page_content.count('\n\n') + 1,
                'has_tables': bool(doc.metadata.get('tables')),
                'has_headers': bool(doc.metadata.get('structure', {}).get('heading')),
            },
            'structural_info': {
                'level': doc.metadata.get('structure', {}).get('level'),
                'heading': doc.metadata.get('structure', {}).get('heading'),
                'is_section_start': bool(doc.metadata.get('structure', {}).get('heading')),
            }
        }
        return metadata